# encoding: utf-8

"""
Test suite for the docx.shape module
"""

from __future__ import absolute_import, print_function, unicode_literals

import pytest

from docx.enum.shape import WD_INLINE_SHAPE
from docx.oxml.document import CT_Body
from docx.oxml.ns import nsmap
from docx.oxml.text.run import CT_R
from docx.parts.document import DocumentPart
from docx.parts.image import ImagePart
from docx.shape import InlineShape, InlineShapes
from docx.shared import Length
from docx.text.run import Run

from .oxml.unitdata.dml import (
    a_blip, a_blipFill, a_cNvGraphicFramePr, a_cNvPr, a_cNvPicPr, a_docPr,
    a_fillRect, a_graphic, a_graphicData, a_graphicFrameLocks, a_pic,
    a_prstGeom, a_stretch, an_ext, an_extent, an_inline, an_nvPicPr, an_off,
    an_spPr, an_xfrm
)
from .oxml.unitdata.text import an_r
from .unitutil.cxml import element, xml
from .unitutil.mock import (
    class_mock, instance_mock, loose_mock, property_mock
)


class DescribeInlineShapes(object):

    def it_knows_how_many_inline_shapes_it_contains(
            self, inline_shapes_fixture):
        inline_shapes, expected_count = inline_shapes_fixture
        assert len(inline_shapes) == expected_count

    def it_can_iterate_over_its_InlineShape_instances(
            self, inline_shapes_fixture):
        inline_shapes, inline_shape_count = inline_shapes_fixture
        actual_count = 0
        for inline_shape in inline_shapes:
            assert isinstance(inline_shape, InlineShape)
            actual_count += 1
        assert actual_count == inline_shape_count

    def it_provides_indexed_access_to_inline_shapes(
            self, inline_shapes_fixture):
        inline_shapes, inline_shape_count = inline_shapes_fixture
        for idx in range(-inline_shape_count, inline_shape_count):
            inline_shape = inline_shapes[idx]
            assert isinstance(inline_shape, InlineShape)

    def it_raises_on_indexed_access_out_of_range(
            self, inline_shapes_fixture):
        inline_shapes, inline_shape_count = inline_shapes_fixture
        with pytest.raises(IndexError):
            too_low = -1 - inline_shape_count
            inline_shapes[too_low]
        with pytest.raises(IndexError):
            too_high = inline_shape_count
            inline_shapes[too_high]

    def it_can_add_an_inline_picture_to_the_document(
            self, add_picture_fixture):
        # fixture ----------------------
        (inline_shapes, image_descriptor_, document_, InlineShape_,
         run, r_, image_part_, rId_, shape_id_, new_picture_shape_
         ) = add_picture_fixture
        # exercise ---------------------
        picture_shape = inline_shapes.add_picture(image_descriptor_, run)
        # verify -----------------------
        document_.get_or_add_image_part.assert_called_once_with(
            image_descriptor_
        )
        InlineShape_.new_picture.assert_called_once_with(
            r_, image_part_, rId_, shape_id_
        )
        assert picture_shape is new_picture_shape_

    def it_knows_the_part_it_belongs_to(self, inline_shapes_with_parent_):
        inline_shapes, parent_ = inline_shapes_with_parent_
        part = inline_shapes.part
        assert part is parent_.part

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def add_picture_fixture(
            self, request, body_, document_, image_descriptor_, InlineShape_,
            r_, image_part_, rId_, shape_id_, new_picture_shape_):
        inline_shapes = InlineShapes(body_, None)
        property_mock(request, InlineShapes, 'part', return_value=document_)
        run = Run(r_, None)
        return (
            inline_shapes, image_descriptor_, document_, InlineShape_, run,
            r_, image_part_, rId_, shape_id_, new_picture_shape_
        )

    @pytest.fixture
    def inline_shapes_fixture(self):
        body = element(
            'w:body/w:p/(w:r/w:drawing/wp:inline, w:r/w:drawing/wp:inline)'
        )
        inline_shapes = InlineShapes(body, None)
        expected_count = 2
        return inline_shapes, expected_count

    # fixture components ---------------------------------------------

    @pytest.fixture
    def body_(self, request, r_):
        body_ = instance_mock(request, CT_Body)
        body_.add_p.return_value.add_r.return_value = r_
        return body_

    @pytest.fixture
    def document_(self, request, rId_, image_part_, shape_id_):
        document_ = instance_mock(request, DocumentPart, name='document_')
        document_.get_or_add_image_part.return_value = image_part_, rId_
        document_.next_id = shape_id_
        return document_

    @pytest.fixture
    def image_part_(self, request):
        return instance_mock(request, ImagePart)

    @pytest.fixture
    def image_descriptor_(self, request):
        return instance_mock(request, str)

    @pytest.fixture
    def InlineShape_(self, request, new_picture_shape_):
        InlineShape_ = class_mock(request, 'docx.shape.InlineShape')
        InlineShape_.new_picture.return_value = new_picture_shape_
        return InlineShape_

    @pytest.fixture
    def inline_shapes_with_parent_(self, request):
        parent_ = loose_mock(request, name='parent_')
        inline_shapes = InlineShapes(None, parent_)
        return inline_shapes, parent_

    @pytest.fixture
    def new_picture_shape_(self, request):
        return instance_mock(request, InlineShape)

    @pytest.fixture
    def r_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def rId_(self, request):
        return instance_mock(request, str)

    @pytest.fixture
    def shape_id_(self, request):
        return instance_mock(request, int)


class DescribeInlineShape(object):

    def it_knows_what_type_of_shape_it_is(self, shape_type_fixture):
        inline_shape, inline_shape_type = shape_type_fixture
        assert inline_shape.type == inline_shape_type

    def it_can_contruct_a_new_inline_picture_shape(
            self, new_picture_fixture):
        inline_shape, r, image_part_, rId, shape_id, expected_inline_xml = (
            new_picture_fixture
        )
        picture = inline_shape.new_picture(r, image_part_, rId, shape_id)
        assert picture._inline.xml == expected_inline_xml
        assert r[0][0] is picture._inline

    def it_knows_its_display_dimensions(self, dimensions_get_fixture):
        inline_shape, cx, cy = dimensions_get_fixture
        width = inline_shape.width
        height = inline_shape.height
        assert isinstance(width, Length)
        assert width == cx
        assert isinstance(height, Length)
        assert height == cy

    def it_can_change_its_display_dimensions(self, dimensions_set_fixture):
        inline_shape, cx, cy, expected_xml = dimensions_set_fixture
        inline_shape.width = cx
        inline_shape.height = cy
        assert inline_shape._inline.xml == expected_xml

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def dimensions_get_fixture(self):
        inline_cxml, expected_cx, expected_cy = (
            'wp:inline/wp:extent{cx=333, cy=666}', 333, 666
        )
        inline_shape = InlineShape(element(inline_cxml))
        return inline_shape, expected_cx, expected_cy

    @pytest.fixture
    def dimensions_set_fixture(self):
        inline_cxml, new_cx, new_cy, expected_cxml = (
            'wp:inline/wp:extent{cx=333, cy=666}', 444, 888,
            'wp:inline/wp:extent{cx=444, cy=888}'
        )
        inline_shape = InlineShape(element(inline_cxml))
        expected_xml = xml(expected_cxml)
        return inline_shape, new_cx, new_cy, expected_xml

    @pytest.fixture
    def new_picture_fixture(self, request, image_part_, image_params):
        filename, rId, cx, cy = image_params
        inline_shape = InlineShape(None)
        r = an_r().with_nsdecls().element
        shape_id = 7
        name = 'Picture %d' % shape_id
        uri = nsmap['pic']
        expected_inline = (
            an_inline().with_nsdecls('wp', 'a', 'pic', 'r', 'w').with_child(
                an_extent().with_cx(cx).with_cy(cy)).with_child(
                a_docPr().with_id(shape_id).with_name(name)).with_child(
                a_cNvGraphicFramePr().with_child(
                    a_graphicFrameLocks().with_noChangeAspect(1))).with_child(
                a_graphic().with_child(
                    a_graphicData().with_uri(uri).with_child(
                        self._pic_bldr(filename, rId, cx, cy))))
        ).element
        expected_inline_xml = expected_inline.xml
        return (
            inline_shape, r, image_part_, rId, shape_id, expected_inline_xml
        )

    @pytest.fixture(params=[
        'embed pic', 'link pic', 'link+embed pic', 'chart', 'smart art',
        'not implemented'
    ])
    def shape_type_fixture(self, request):
        if request.param == 'embed pic':
            inline = self._inline_with_picture(embed=True)
            shape_type = WD_INLINE_SHAPE.PICTURE

        elif request.param == 'link pic':
            inline = self._inline_with_picture(link=True)
            shape_type = WD_INLINE_SHAPE.LINKED_PICTURE

        elif request.param == 'link+embed pic':
            inline = self._inline_with_picture(embed=True, link=True)
            shape_type = WD_INLINE_SHAPE.LINKED_PICTURE

        elif request.param == 'chart':
            inline = self._inline_with_uri(nsmap['c'])
            shape_type = WD_INLINE_SHAPE.CHART

        elif request.param == 'smart art':
            inline = self._inline_with_uri(nsmap['dgm'])
            shape_type = WD_INLINE_SHAPE.SMART_ART

        elif request.param == 'not implemented':
            inline = self._inline_with_uri('foobar')
            shape_type = WD_INLINE_SHAPE.NOT_IMPLEMENTED

        return InlineShape(inline), shape_type

    # fixture components ---------------------------------------------

    @pytest.fixture
    def image_params(self):
        filename = 'foobar.garf'
        rId = 'rId42'
        cx, cy = 914422, 223344
        return filename, rId, cx, cy

    @pytest.fixture
    def image_part_(self, request, image_params):
        filename, rId, cx, cy = image_params
        image_part_ = instance_mock(request, ImagePart)
        image_part_.default_cx = cx
        image_part_.default_cy = cy
        image_part_.filename = filename
        return image_part_

    def _inline_with_picture(self, embed=False, link=False):
        picture_ns = nsmap['pic']

        blip_bldr = a_blip()
        if embed:
            blip_bldr.with_embed('rId1')
        if link:
            blip_bldr.with_link('rId2')

        inline = (
            an_inline().with_nsdecls('wp', 'r').with_child(
                a_graphic().with_nsdecls().with_child(
                    a_graphicData().with_uri(picture_ns).with_child(
                        a_pic().with_nsdecls().with_child(
                            a_blipFill().with_child(
                                blip_bldr)))))
        ).element
        return inline

    def _inline_with_uri(self, uri):
        inline = (
            an_inline().with_nsdecls('wp').with_child(
                a_graphic().with_nsdecls().with_child(
                    a_graphicData().with_uri(uri)))
        ).element
        return inline

    def _pic_bldr(self, name, rId, cx, cy):
        return (
            a_pic().with_child(
                an_nvPicPr().with_child(
                    a_cNvPr().with_id(0).with_name(name)).with_child(
                    a_cNvPicPr())).with_child(
                a_blipFill().with_child(
                    a_blip().with_embed(rId)).with_child(
                    a_stretch().with_child(
                        a_fillRect()))).with_child(
                an_spPr().with_child(
                    an_xfrm().with_child(
                        an_off().with_x(0).with_y(0)).with_child(
                        an_ext().with_cx(cx).with_cy(cy))).with_child(
                    a_prstGeom().with_prst('rect')))
        )
